import os
import json
import datetime
import traceback
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CORRECCIONES DE IMPORTACIÓN ---
from kivy.utils import platform 
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.scrollview import ScrollView
from kivy.uix.screenmanager import ScreenManager, Screen, FadeTransition
from kivy.uix.textinput import TextInput
from kivy.metrics import dp
from kivy.core.window import Window

# --- RUTAS DE ALMACENAMIENTO CORREGIDAS ---
if platform == 'android':
    try:
        from android.storage import primary_external_storage_path
        base_dir = primary_external_storage_path()
        SAVE_PATH = os.path.join(base_dir, "Download", "YOURMOBILE_CAJA")
    except:
        SAVE_PATH = "/sdcard/Download/YOURMOBILE_CAJA"
else:
    SAVE_PATH = "YOURMOBILE_CAJA"

STAFF_FILE = os.path.join(SAVE_PATH, "staff_list.txt")
TEMP_FILE = os.path.join(SAVE_PATH, "daily_temp_cache.json")
LOG_FILE = os.path.join(SAVE_PATH, "ERROR_LOG_YOURMOBILE.txt")

try:
    if not os.path.exists(SAVE_PATH):
        os.makedirs(SAVE_PATH, exist_ok=True)
except Exception as e:
    print(f"Error creando carpeta: {e}")

class MenuPrincipal(Screen):
    def on_pre_enter(self):
        self.lbl_info.text = f"Items in daily memory: {len(App.get_running_app().products)}"

    def __init__(self, **kw):
        super().__init__(**kw)
        layout = BoxLayout(orientation='vertical', padding=dp(10), spacing=dp(5))
        layout.add_widget(Label(text="YOUR MOBILE STORE", font_size='24sp', bold=True, size_hint_y=None, height=dp(50), color=(0.29, 0.56, 0.88, 1)))
        
        scroll = ScrollView()
        self.btn_container = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(6))
        self.btn_container.bind(minimum_height=self.btn_container.setter('height'))

        # SECCIONES COMPLETAS RECTIFICADAS
        secciones = [
            ("--- INCOMES ---", [("[ 1 ] NEW SALE (PRODUCT)", "PRODUCT"), ("[ 2 ] REPAIR / SERVICE", "SERVICE/REPAIR")]),
            ("--- INVESTMENTS & TAXI ---", [
                ("[ 3 ] INVESTMENT SAN JUAN", "INVESTMENT SAN JUAN"), 
                ("[ 4 ] INVESTMENT TOBAGO", "INVESTMENT TOBAGO"),
                ("[ 5 ] INVESTMENT ARANGUEZ", "INVESTMENT ARANGUEZ")
            ]),
            ("--- OTHER OUTFLOWS ---", [
                ("[ 6 ] OWNER'S SON", "OWNER'S SON"), 
                ("[ 7 ] VOUCHER", "VOUCHER"), 
                ("[ 8 ] RENT PAYMENT", "RENT PAYMENT"), 
                ("[ 9 ] WAGES / SALARIES", "WAGES")
            ]),
        ]

        for titulo, botones in secciones:
            self.btn_container.add_widget(Label(text=titulo, size_hint_y=None, height=dp(35), color=(0.4, 0.4, 0.4, 1), bold=True))
            for texto, cat in botones:
                btn = Button(text=texto, size_hint_y=None, height=dp(50), background_color=(0.15, 0.5, 0.3, 1), halign='left', padding=(dp(15), 0))
                btn.bind(on_release=lambda x, c=cat: self.ir_a_flujo(c))
                btn.text_size = (Window.width - dp(40), None)
                self.btn_container.add_widget(btn)

        self.btn_container.add_widget(Label(text="--- MANAGEMENT ---", size_hint_y=None, height=dp(35), color=(0.4, 0.4, 0.4, 1), bold=True))
        btn_staff = Button(text="[ 10 ] STAFF MANAGEMENT", size_hint_y=None, height=dp(55), background_color=(0.16, 0.5, 0.72, 1), bold=True)
        btn_staff.bind(on_release=lambda x: setattr(self.manager, 'current', 'staff'))
        self.btn_container.add_widget(btn_staff)

        scroll.add_widget(self.btn_container)
        layout.add_widget(scroll)
        
        btn_save = Button(text=" SAVE REPORT AND EXIT", size_hint_y=None, height=dp(60), background_color=(0.11, 0.51, 0.28, 1), bold=True)
        btn_save.bind(on_press=lambda x: App.get_running_app().generate_report())
        layout.add_widget(btn_save)

        self.lbl_info = Label(text="Items: 0", size_hint_y=None, height=dp(30), font_size='10sp')
        layout.add_widget(self.lbl_info)
        self.add_widget(layout)

    def ir_a_flujo(self, cat):
        self.manager.get_screen('input').category = cat
        self.manager.current = 'input'

class StaffScreen(Screen):
    def on_pre_enter(self): self.refresh_list()
    def __init__(self, **kw):
        super().__init__(**kw)
        layout = BoxLayout(orientation='vertical', padding=dp(20), spacing=dp(10))
        header = BoxLayout(size_hint_y=None, height=dp(50))
        btn_back = Button(text="‹ BACK", size_hint_x=0.2, background_color=(0.3, 0.3, 0.3, 1))
        btn_back.bind(on_release=lambda x: setattr(self.manager, 'current', 'menu'))
        header.add_widget(btn_back); header.add_widget(Label(text="STAFF MANAGEMENT", bold=True))
        layout.add_widget(header)
        
        add_box = BoxLayout(size_hint_y=None, height=dp(50), spacing=dp(5))
        self.new_staff_input = TextInput(hint_text="Worker Name", multiline=False)
        btn_add = Button(text="ADD", size_hint_x=0.3, background_color=(0.15, 0.68, 0.37, 1))
        btn_add.bind(on_press=self.add_worker)
        add_box.add_widget(self.new_staff_input); add_box.add_widget(btn_add); layout.add_widget(add_box)
        
        self.scroll = ScrollView(); self.list_layout = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(5))
        self.list_layout.bind(minimum_height=self.list_layout.setter('height')); self.scroll.add_widget(self.list_layout); layout.add_widget(self.scroll)
        self.add_widget(layout)

    def refresh_list(self):
        self.list_layout.clear_widgets()
        if os.path.exists(STAFF_FILE):
            with open(STAFF_FILE, "r") as f:
                for line in f:
                    name = line.strip()
                    if name:
                        row = BoxLayout(size_hint_y=None, height=dp(45), spacing=dp(5))
                        row.add_widget(Label(text=name, halign='left'))
                        btn_del = Button(text="DEL", size_hint_x=0.2, background_color=(0.75, 0.22, 0.16, 1))
                        btn_del.bind(on_release=lambda x, n=name: self.delete_worker(n))
                        row.add_widget(btn_del); self.list_layout.add_widget(row)

    def add_worker(self, instance):
        name = self.new_staff_input.text.strip().upper()
        if name:
            with open(STAFF_FILE, "a") as f: f.write(name + "\n")
            self.new_staff_input.text = ""; self.refresh_list()

    def delete_worker(self, name):
        if os.path.exists(STAFF_FILE):
            with open(STAFF_FILE, "r") as f: lines = f.readlines()
            with open(STAFF_FILE, "w") as f:
                for line in lines:
                    if line.strip() != name: f.write(line)
            self.refresh_list()

class InputScreen(Screen):
    category = ""
    fields = {}

    def on_pre_enter(self):
        self.layout_campos.clear_widgets(); self.fields = {}
        self.lbl_header.text = self.category
        if self.category == "WAGES":
            self.create_field('worker', "WORKER NAME:"); self.create_field('amt', "SALARY AMOUNT:")
        elif "INVESTMENT" in self.category:
            self.create_field('amt', "INVESTED AMOUNT:"); self.create_field('taxi', "TAXI COST:")
        elif self.category == "RENT PAYMENT":
            self.create_field('aranguez', "RENT ARANGUEZ (SHOP):")
            self.create_field('sanjuan', "RENT SAN JUAN (SHOP):")
            self.create_field('home', "RENT HOME (PERSONAL):")
        elif self.category in ["OWNER'S SON", "VOUCHER"]:
            self.create_field('amt', "AMOUNT PAID:")
        else:
            self.create_field('name', "ITEM NAME:"); self.create_field('qty', "QTY:")
            self.create_field('price', "SALE PRICE:"); self.create_field('cost', "COST PRICE:"); self.create_field('taxi', "TAXI FEE:")

    def create_field(self, key, label_text):
        self.layout_campos.add_widget(Label(text=label_text, size_hint_y=None, height=dp(25), color=(0.29, 0.56, 0.88, 1)))
        ti = TextInput(multiline=False, size_hint_y=None, height=dp(45))
        self.fields[key] = ti; self.layout_campos.add_widget(ti)

    def process_data(self, instance):
        app = App.get_running_app(); f = self.fields
        def n(k): 
            try: return float(f[k].text) if f[k].text else 0.0
            except: return 0.0

        if self.category == "WAGES":
            amt = n('amt')
            app.products.append({"name": f"WAGE: {f['worker'].text.upper()}", "price": 0.0, "cost": amt, "transport": 0.0, "quantity": 1, "total": -amt})
        elif "INVESTMENT" in self.category:
            amt, taxi = n('amt'), n('taxi')
            app.products.append({"name": self.category, "price": 0.0, "cost": amt, "transport": 0.0, "quantity": 1, "total": -amt})
            app.products.append({"name": f"TAXI: {self.category}", "price": 0.0, "cost": 0.0, "transport": taxi, "quantity": 1, "total": -taxi})
        elif self.category == "RENT PAYMENT":
            for k, key in [("ARANGUEZ (SHOP)", 'aranguez'), ("SAN JUAN (SHOP)", 'sanjuan'), ("HOME (PERSONAL)", 'home')]:
                val = n(key)
                if val > 0: app.products.append({"name": f"RENT: {k}", "price": 0.0, "cost": val, "transport": 0.0, "quantity": 1, "total": -val})
        elif self.category in ["OWNER'S SON", "VOUCHER"]:
            amt = n('amt')
            app.products.append({"name": f"OUTFLOW: {self.category}", "price": 0.0, "cost": amt, "transport": 0.0, "quantity": 1, "total": -amt})
        else:
            name, qty = f['name'].text.upper() or "ITEM", int(n('qty')) or 1
            ps, cb, tr = n('price'), n('cost'), n('taxi')
            net = (ps * qty) - (cb * qty) - tr
            app.products.append({"name": name if self.category != "SERVICE/REPAIR" else f"SRV: {name}", "price": ps * qty, "cost": cb * qty, "transport": tr, "quantity": qty, "total": net})

        app.save_temporary_data(); self.manager.current = 'menu'

    def __init__(self, **kw):
        super().__init__(**kw)
        layout = BoxLayout(orientation='vertical', padding=dp(20), spacing=dp(10))
        header = BoxLayout(size_hint_y=None, height=dp(50))
        btn_back = Button(text="‹ BACK", size_hint_x=0.25, background_color=(0.3, 0.3, 0.3, 1))
        btn_back.bind(on_release=lambda x: setattr(self.manager, 'current', 'menu'))
        self.lbl_header = Label(text="", bold=True)
        header.add_widget(btn_back); header.add_widget(self.lbl_header)
        self.layout_campos = BoxLayout(orientation='vertical', spacing=dp(5))
        btn_confirm = Button(text="CONFIRM RECORD", size_hint_y=None, height=dp(60), background_color=(0.15, 0.68, 0.37, 1), bold=True)
        btn_confirm.bind(on_press=self.process_data)
        layout.add_widget(header); layout.add_widget(self.layout_campos); layout.add_widget(BoxLayout()); layout.add_widget(btn_confirm)
        self.add_widget(layout)

class MobileStoreApp(App):
    products = []
    def build(self):
        self.load_temporary_data()
        sm = ScreenManager(transition=FadeTransition())
        sm.add_widget(MenuPrincipal(name='menu')); sm.add_widget(InputScreen(name='input')); sm.add_widget(StaffScreen(name='staff'))
        return sm
    def load_temporary_data(self):
        if os.path.exists(TEMP_FILE):
            try:
                with open(TEMP_FILE, "r") as f: self.products = json.load(f)
            except: self.products = []
    def save_temporary_data(self):
        try:
            with open(TEMP_FILE, "w") as f: json.dump(self.products, f)
        except Exception as e:
            with open(LOG_FILE, "a") as log: log.write(f"Error saving temp: {e}\n")

    def generate_report(self):
        if not self.products: self.stop(); return
        try:
            doc = Document(); doc.add_heading("YOUR MOBILE STORE", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            now = datetime.datetime.now(); doc.add_heading(f"REPORT: {now.strftime('%m/%d/%Y %I:%M %p')}", level=1)
            table = doc.add_table(rows=1, cols=6); table.style = 'Table Grid'
            for i, t in enumerate(['DESC', 'QTY', 'SALE', 'COST', 'FEE', 'NET']): table.rows[0].cells[i].text = t
            for p in self.products:
                r = table.add_row().cells
                r[0].text, r[1].text, r[2].text, r[3].text, r[4].text, r[5].text = str(p['name']), str(p['quantity']), f"{p['price']:.2f}", f"{p['cost']:.2f}", f"{p['transport']:.2f}", f"{p['total']:.2f}"
            doc.save(os.path.join(SAVE_PATH, f"REPORT_{now.strftime('%Y%m%d_%H%M')}.docx"))
            if os.path.exists(TEMP_FILE): os.remove(TEMP_FILE)
            self.stop()
        except Exception as e:
            with open(LOG_FILE, "a") as log: log.write(f"Error docx: {e}\n")

if __name__ == '__main__':
    try:
        MobileStoreApp().run()
    except Exception:
        with open(LOG_FILE, "w") as f:
            f.write(traceback.format_exc())
        
